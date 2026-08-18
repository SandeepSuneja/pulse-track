"""Generate docs/Pulse-Track-AWS-Deployment-and-Architecture.docx"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm, Emu

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Pulse-Track-AWS-Deployment-and-Architecture.docx"
DIAGRAM = ROOT / "architecture-aws.png"

NAVY = RGBColor(0x0F, 0x2C, 0x59)
ACCENT = RGBColor(0x1D, 0x4E, 0x89)
GRAY = RGBColor(0x44, 0x44, 0x44)
RULE = "D6DEE8"


def set_run_font(run, *, size=11, bold=False, italic=False, color=GRAY, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C5CDD6")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pulse Track  ·  AWS Deployment & Architecture  ·  Page ")
    set_run_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66))
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = p.add_run()
    r._r.append(fld)
    r._r.append(instr)
    r._r.append(fld2)
    set_run_font(r, size=9, color=RGBColor(0x66, 0x66, 0x66))


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else ACCENT
        run.font.name = "Calibri"
    return p


def para(doc, text, *, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=GRAY)
    return p


def bullets(doc, items, *, bold_lead=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if bold_lead and " — " in item:
            lead, rest = item.split(" — ", 1)
            r1 = p.add_run(lead)
            set_run_font(r1, bold=True, color=NAVY)
            r2 = p.add_run(" — " + rest)
            set_run_font(r2, color=GRAY)
        else:
            r = p.add_run(item)
            set_run_font(r, color=GRAY)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, color=GRAY)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text.strip("\n"))
    set_run_font(run, size=9, name="Consolas", color=RGBColor(0x1A, 0x1A, 0x1A))
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F7FA")
    shd.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "0F2C59")
        set_cell_border(hdr[i])
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        fill = "FFFFFF" if ri % 2 == 0 else "F4F7FA"
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=10, color=GRAY)
            shade_cell(cells[ci], fill)
            set_cell_border(cells[ci])
    doc.add_paragraph()


def qa(doc, q, a):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q. " + q)
    set_run_font(r, size=11, bold=True, color=NAVY)
    para(doc, "A. " + a, size=11, space_after=6)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_page_number(section)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.paragraph_format.space_before = Pt(48)
    r = t.add_run("PULSE TRACK")
    set_run_font(r, size=14, bold=True, color=ACCENT)

    t = doc.add_paragraph()
    r = t.add_run("AWS Production Deployment\nArchitecture & Interview Guide")
    set_run_font(r, size=28, bold=True, color=NAVY)

    para(
        doc,
        "A reusable runbook for deploying Pulse Track on AWS, with the full architecture, "
        "why each service exists, how traffic and secrets flow, how to copy local data, "
        "and how to explain the design in interviews.",
        size=12,
    )
    table(
        doc,
        ["Field", "Value"],
        [
            ["Document type", "Architecture + deployment runbook + interview prep"],
            ["Audience", "You (future deploys) and interviewers discussing this project"],
            ["Application", "Pulse Track — personal board, time logs, goals, analytics"],
            ["Stack", "React (Vite) SPA + FastAPI + Firebase Auth + PostgreSQL"],
            ["Cloud pattern", "Static frontend + containerized API + managed database"],
            ["Live region (this account)", "us-east-1 (N. Virginia)"],
            ["Original plan region", "ap-south-1 is equally valid; swap the region in every command"],
            ["Version", "August 2026"],
        ],
    )

    heading(doc, "1. Purpose of this document")
    para(
        doc,
        "This file is the single place to re-create the production environment, understand "
        "why the pieces were chosen, and talk through the system under interview pressure. "
        "It is not a dump of AWS console screenshots. It is the design plus the exact order "
        "of operations."
    )
    bullets(
        doc,
        [
            "Later deploys — follow sections 8–12 in order. Commands are PowerShell.",
            "Interviews — use sections 3–7 and 14. Those explain trade-offs, not just names of services.",
            "Secrets — never paste DATABASE_URL, Firebase private keys, or IAM access keys into this document or Git.",
        ],
        bold_lead=True,
    )

    heading(doc, "2. What Pulse Track is")
    para(
        doc,
        "Pulse Track is a personal workspace: a Kanban board, time logs against in-progress "
        "tasks, goals (hours target or a deadline), and charts. Users sign in with Google or "
        "email/password. The product is split into two apps that share one API and one database "
        "so a future mobile client can reuse everything."
    )
    table(
        doc,
        ["Area", "What it stores / shows"],
        [
            ["Board", "Tasks: todo → in_progress → completed; dates, category, optional goal"],
            ["Activities", "Time logs against in-progress tasks"],
            ["Goals", "Hours target or deadline; complete / failed"],
            ["Dashboard / Analytics", "Minutes by period, category, and task"],
            ["Profile", "Display name, bio, timezone"],
        ],
    )
    para(doc, "REST resources the API exposes (all under /api, all user-scoped after auth):")
    bullets(
        doc,
        [
            "GET /api/health — liveness probe (no auth)",
            "GET/PATCH /api/users/me",
            "GET/POST /api/tasks, GET/PATCH/DELETE /api/tasks/{id}",
            "GET/POST /api/activities, GET/PATCH/DELETE /api/activities/{id}",
            "GET/POST /api/goals, PATCH/DELETE /api/goals/{id}",
            "GET /api/analytics/summary?period=day|week|month|year",
        ],
    )

    heading(doc, "3. Design goals (say this first in interviews)")
    numbered(
        doc,
        [
            "One public HTTPS API for web and mobile. Do not bury the API behind the SPA host.",
            "No passwords in this app. Firebase issues ID tokens; FastAPI verifies them.",
            "Durable shared data. SQLite is a local file; production needs a real database server.",
            "Secrets out of Git and out of the Docker image. Runtime config from Parameter Store.",
            "Least privilege on the network. The database is not on the public internet.",
            "Cheap operations. No EC2 patching; containers on Fargate; static files on S3.",
        ],
    )

    heading(doc, "4. Architecture overview")
    para(
        doc,
        "Production is three planes: a static web plane (S3 + CloudFront), an API plane "
        "(ALB + ECS Fargate), and a data plane (RDS PostgreSQL). Identity is Firebase. "
        "Config is SSM Parameter Store. Delivery of new versions is GitHub Actions → ECR / S3."
    )
    if DIAGRAM.exists():
        cap = doc.add_paragraph()
        r = cap.add_run("Figure 1. Pulse Track on AWS — clients, edge, VPC, data, identity, CI/CD.")
        set_run_font(r, size=10, italic=True, color=ACCENT)
        doc.add_picture(str(DIAGRAM), width=Inches(6.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "4.1 Request flows", 2)
    para(doc, "Web page load", bold=True)
    numbered(
        doc,
        [
            "Browser requests https://<cloudfront-domain>/ (and later /board, /goals, …).",
            "CloudFront serves index.html and hashed JS/CSS from the private S3 bucket via Origin Access Control.",
            "Because this is a React Router SPA, unknown paths are rewritten to /index.html (CloudFront custom errors 403/404 → 200 /index.html).",
            "The JS bundle already contains VITE_API_URL (baked in at npm run build) and Firebase web config.",
        ],
    )
    para(doc, "Authenticated API call (web or mobile)", bold=True)
    numbered(
        doc,
        [
            "User signs in with Firebase (Google popup or email/password).",
            "Client holds a Firebase ID token and sends Authorization: Bearer <token> on every API request.",
            "Request hits the public HTTPS endpoint (ECS Express URL or an ALB DNS name).",
            "ALB forwards to a Fargate task running uvicorn on port 8000.",
            "FastAPI verifies the token with Firebase Admin SDK, then loads or creates a users row by firebase_uid.",
            "SQLAlchemy talks to RDS PostgreSQL on 5432 inside the VPC. Response JSON goes back to the client.",
        ],
    )
    para(doc, "Mobile (when you build it)", bold=True)
    para(
        doc,
        "The native app does not use CloudFront or S3. It uses the same API origin, the same "
        "Firebase project, and the same Bearer header. CORS is a browser rule; native HTTP "
        "clients ignore it. Only a WebView would need an extra origin in CORS_ORIGINS."
    )

    heading(doc, "5. Component encyclopedia")
    para(
        doc,
        "For each piece: what it is, why it was chosen, and why the system actually needs it. "
        "Interviewers care more about the “why” than the logo."
    )

    components = [
        (
            "React + Vite SPA",
            "The UI. Vite emits static files into frontend/dist/.",
            "A SPA can be hosted without a Node server. That lets us put the UI on S3+CloudFront and keep the API independently scalable.",
            "Necessary as the product UI. Necessary as a static artifact so web hosting stays cheap and cacheable.",
        ),
        (
            "FastAPI + Uvicorn",
            "The HTTP API (app.main:app) on port 8000.",
            "Typed REST, OpenAPI at /docs, SQLAlchemy, easy to containerize. One process serves web and future mobile.",
            "Necessary as the only place that talks to the database and enforces per-user scoping.",
        ),
        (
            "Firebase Authentication",
            "Google and email/password on the client; Admin SDK verify_id_token on the server.",
            "We do not store passwords. Cognito would also work; Firebase was already in the app and is enough for web+mobile.",
            "Necessary for identity. Without it anyone could call /api/tasks. DEV_SKIP_AUTH is local-only and must stay false in production.",
        ),
        (
            "Amazon S3",
            "Private bucket holding the built SPA (index.html, hashed assets).",
            "Object storage is the native home for static files. No OS to patch. Durability is 11 nines.",
            "Necessary as the origin for CloudFront. The bucket stays private; the public internet never lists it.",
        ),
        (
            "Amazon CloudFront",
            "HTTPS CDN in front of S3, with OAC and SPA error pages.",
            "Gives a free HTTPS hostname, caches JS/CSS at edge, hides S3, and rewrites SPA routes. Putting the bucket public would skip HTTPS customization and leak listing.",
            "Necessary for HTTPS web access and for /board deep links. Without the 403/404 → index.html mapping, refresh on /board returns S3’s XML error.",
        ),
        (
            "CloudFront Origin Access Control (OAC)",
            "CloudFront signs requests to S3; bucket policy allows only that distribution.",
            "Replaces the older OAI. Stops people from bypassing CloudFront and hitting the bucket URL.",
            "Necessary for a private bucket that is still readable by the CDN.",
        ),
        (
            "Docker + Amazon ECR",
            "backend/Dockerfile builds a Python 3.12 image; ECR stores pulse-track-api.",
            "Same artifact locally, in CI, and on Fargate. ECR is private and in-region, so pulls are fast and authenticated by IAM.",
            "Necessary because ECS runs containers, not a folder of .py files. Never COPY .env or firebase-service-account.json into the image (.dockerignore).",
        ),
        (
            "Amazon ECS on AWS Fargate",
            "Runs the API container without managing EC2.",
            "App Runner was the first plan (simpler PaaS). AWS closed App Runner to new customers on 30 April 2026. ECS Express Mode is the official replacement: one API call provisions Fargate + ALB + autoscaling + HTTPS URL. You still get a real ECS service you can update from GitHub Actions.",
            "Necessary to keep the API process running, health-checked, and replaceable. Fargate removes instance patching, which is the usual interview answer for “why not EC2 for this size of app”.",
        ),
        (
            "Application Load Balancer",
            "Public HTTPS listener; targets the Fargate tasks on 8000; health check GET /api/health.",
            "Gives a stable hostname, TLS termination, and the ability to run 2+ tasks later. ECS Express Mode creates this for you.",
            "Necessary as the internet door into the VPC. Tasks themselves should not have public IPs for the database path; the ALB is the front door.",
        ),
        (
            "Amazon RDS for PostgreSQL 16",
            "Managed database pulse-track-db, db.t4g.micro, 20 GiB gp3, not publicly accessible.",
            "SQLite cannot be shared by multiple API tasks or by web+mobile safely (file locking, no network). RDS handles backups, patching, and a stable endpoint. SQLAlchemy already talks to Postgres via psycopg2; sqlite PRAGMA migrations are skipped when DATABASE_URL is not sqlite.",
            "Necessary for durable, concurrent, shared state. Tables are created on API startup (Base.metadata.create_all).",
        ),
        (
            "VPC + subnets",
            "Default VPC is enough for a first deploy. RDS lives on subnets; ECS tasks use a security group in the same VPC.",
            "A VPC is the blast-radius boundary. Private RDS means no 5432 from the internet.",
            "Necessary so the API can reach RDS without making the database public. Temporarily opening public access is only for one-off laptop imports, then it must be turned off.",
        ),
        (
            "Security groups",
            "pulse-track-rds-sg allows TCP 5432 only from the API task SG. ALB SG allows 443 from the internet. Task SG allows 8000 from the ALB SG.",
            "Stateful firewalls that encode least privilege. “SG referencing SG” is better than opening 5432 to 0.0.0.0/0.",
            "Necessary. A public RDS with 5432 open to the world is the most common way this architecture gets failed in a security review.",
        ),
        (
            "AWS Systems Manager Parameter Store",
            "Names under /pulse-track/: DATABASE_URL and FIREBASE_*.",
            "Encrypted at rest (use SecureString). ECS task execution role calls ssm:GetParameters and injects env vars. Unlike baking secrets into the image, you can rotate a password without rebuilding Docker.",
            "Necessary so production config is not in Git. Prefer SecureString over String. The live account currently has String type on some parameters — tightening that is a good follow-up.",
        ),
        (
            "IAM roles",
            "ecsTaskExecutionRole (ECR pull, CloudWatch logs, SSM GetParameters). Infrastructure role for Express Mode. Separate GitHub user for CI.",
            "Roles, not long-lived keys on the server. GitHub uses a dedicated user; do not reuse the console admin user in Actions.",
            "Necessary for every AWS API call the platform makes on your behalf. iam:PassRole must list the execution role or CI cannot register a new task definition.",
        ),
        (
            "GitHub Actions",
            ".github/workflows/deploy.yml on push to master: build/push ECR, render task definition, update ECS; optional S3 sync + CloudFront invalidation.",
            "Repeatable deploys. Manual docker push is a fallback. Does not run on feature branches unless you change the workflow.",
            "Necessary for a second deploy without remembering console clicks. First-time AWS resources (RDS, cluster, bucket) are still created once by hand.",
        ),
        (
            "Health check /api/health",
            "Returns {\"status\":\"ok\",\"service\":\"pulse-track\"}.",
            "ALB and ECS use it to stop sending traffic to a dead task (bad DATABASE_URL, Firebase init failure, crash).",
            "Necessary. Without it, a broken container can stay in service and look “green” in the console.",
        ),
    ]
    for name, what, why, need in components:
        heading(doc, name, 3)
        para(doc, "What it is. " + what)
        para(doc, "Why we used it. " + why)
        para(doc, "Why it is necessary. " + need)

    heading(doc, "6. Local versus production")
    table(
        doc,
        ["Concern", "Laptop", "AWS"],
        [
            ["UI", "npm run dev :5173", "S3 + CloudFront HTTPS"],
            ["API", "uvicorn --reload :8000", "Fargate + ALB HTTPS :443 → :8000"],
            ["Database", "SQLite file backend/pulse_track.db", "RDS PostgreSQL pulsetrack"],
            ["Auth skip", "DEV_SKIP_AUTH=true allowed", "Must be false"],
            ["Firebase Admin", "JSON file or .env", "SSM → container env"],
            ["CORS", "http://localhost:5173", "https://<cloudfront-domain>"],
            ["API URL in UI", "http://127.0.0.1:8000", "VITE_API_URL at build time"],
            ["How many API copies", "One process", "1–N tasks behind a load balancer"],
        ],
    )
    para(
        doc,
        "Interview line: “We kept SQLite for zero-ops local development and swapped the "
        "SQLAlchemy URL in production. The schema create path is the same; only SQLite-specific "
        "PRAGMA migrations are skipped on Postgres.”"
    )

    heading(doc, "7. Security model")
    heading(doc, "7.1 Authentication and authorization", 2)
    numbered(
        doc,
        [
            "Firebase authenticates the human (Google or password).",
            "The client never sends a database password. It sends a short-lived ID token.",
            "FastAPI is the PEP (policy enforcement point): invalid token → 401; missing user row → upsert by firebase_uid.",
            "Every task/activity/goal query is filtered by current_user.id. That is authorization (tenancy), not just authentication.",
        ],
    )
    heading(doc, "7.2 Secrets", 2)
    bullets(
        doc,
        [
            "Git ignores .env, *.db, firebase-service-account.json.",
            "Docker ignores the same files.",
            "Production reads FIREBASE_* and DATABASE_URL from SSM.",
            "CORS_ORIGINS and DEV_SKIP_AUTH are non-secret env vars on the task.",
        ],
    )
    heading(doc, "7.3 Network", 2)
    bullets(
        doc,
        [
            "S3: Block Public Access + OAC-only bucket policy.",
            "RDS: PubliclyAccessible = false; 5432 only from the API security group.",
            "API: only 443 on the ALB is public.",
            "If you open RDS to your laptop to import data, revoke the CIDR and set PubliclyAccessible = false the same day.",
        ],
    )
    heading(doc, "7.4 CORS", 2)
    para(
        doc,
        "CORS is a browser sandbox. The API allow-list must be exactly the CloudFront origin "
        "(https://dxxxx.cloudfront.net, no trailing slash). Mobile native apps are not browsers. "
        "A mismatch here is the usual “it works in curl but not in Chrome” bug."
    )

    heading(doc, "8. Repeatable deployment runbook")
    para(
        doc,
        "Do this on a new AWS account or a new region. Replace account IDs and names. "
        "The live account used us-east-1 after App Runner was unavailable; either region works."
    )

    heading(doc, "8.0 Prerequisites", 2)
    bullets(
        doc,
        [
            "AWS CLI v2 configured (aws sts get-caller-identity).",
            "Docker Desktop running.",
            "Node 22+ and Python 3.10+ for local builds.",
            "Firebase project with Google + Email/Password enabled, web app config, and a service-account JSON (keep it off Git).",
            "IAM principal that can manage VPC, RDS, ECR, ECS, S3, CloudFront, SSM, IAM.",
        ],
    )
    code_block(
        doc,
        """$Region = "us-east-1"
$AccountId = (aws sts get-caller-identity --query Account --output text)
$EcrUri = "$AccountId.dkr.ecr.$Region.amazonaws.com"
$ApiRepo = "pulse-track-api"
$WebBucket = "pulse-track-web-$AccountId"
""",
    )

    heading(doc, "8.1 Network and RDS", 2)
    numbered(
        doc,
        [
            "Note the default VPC ID and at least two subnet IDs in different AZs.",
            "Create security group pulse-track-api-sg (you can keep an older name such as pulse-track-apprunner-sg). No inbound required for RDS; ALB/Express Mode will add 443 as needed.",
            "Create pulse-track-rds-sg. Inbound: PostgreSQL 5432 sourced from pulse-track-api-sg only.",
            "Create RDS PostgreSQL 16, identifier pulse-track-db, db.t4g.micro, 20 GiB gp3, public access No, SG = pulse-track-rds-sg, initial DB name pulsetrack. Save the master password.",
            "Connection string shape: postgresql+psycopg2://USER:PASSWORD@ENDPOINT:5432/pulsetrack (URL-encode special characters in the password).",
        ],
    )

    heading(doc, "8.2 SSM parameters", 2)
    para(
        doc,
        "Create these names (SecureString). Values come from the Firebase service-account JSON "
        "plus the RDS URL. The private_key must keep \\n sequences as two characters so FastAPI "
        "can expand them to real newlines."
    )
    table(
        doc,
        ["SSM name", "Maps to container env"],
        [
            ["/pulse-track/DATABASE_URL", "DATABASE_URL"],
            ["/pulse-track/FIREBASE_PROJECT_ID", "FIREBASE_PROJECT_ID"],
            ["/pulse-track/FIREBASE_PRIVATE_KEY_ID", "FIREBASE_PRIVATE_KEY_ID"],
            ["/pulse-track/FIREBASE_PRIVATE_KEY", "FIREBASE_PRIVATE_KEY"],
            ["/pulse-track/FIREBASE_CLIENT_EMAIL", "FIREBASE_CLIENT_EMAIL"],
            ["/pulse-track/FIREBASE_CLIENT_ID", "FIREBASE_CLIENT_ID"],
            ["/pulse-track/FIREBASE_CLIENT_CERT_URL", "FIREBASE_CLIENT_CERT_URL"],
        ],
    )

    heading(doc, "8.3 IAM for ECS", 2)
    bullets(
        doc,
        [
            "ecsTaskExecutionRole + AmazonECSTaskExecutionRolePolicy.",
            "Inline ssm:GetParameters on arn:aws:ssm:REGION:ACCOUNT:parameter/pulse-track/*.",
            "ecsInfrastructureRoleForExpressServices — let the console create it if prompted.",
        ],
    )

    heading(doc, "8.4 Build and push the API image", 2)
    code_block(
        doc,
        """aws ecr create-repository --repository-name $ApiRepo --region $Region
aws ecr get-login-password --region $Region | docker login --username AWS --password-stdin $EcrUri
docker build -t $ApiRepo ./backend
docker tag "${ApiRepo}:latest" "${EcrUri}/${ApiRepo}:latest"
docker push "${EcrUri}/${ApiRepo}:latest"
""",
    )

    heading(doc, "8.5 Create the ECS Express Mode service", 2)
    para(
        doc,
        "Console: ECS → Express Mode → Create. Or AWS CLI create-express-gateway-service. "
        "This replaces App Runner for new accounts."
    )
    table(
        doc,
        ["Setting", "Value"],
        [
            ["Image", "<account>.dkr.ecr.<region>.amazonaws.com/pulse-track-api:latest"],
            ["Service name", "pulse-track-api"],
            ["Container port", "8000"],
            ["Health check path", "/api/health"],
            ["CPU / memory", "1 vCPU / 2 GB is enough"],
            ["Min / max tasks", "1 / 4"],
            ["Plain env", "DEV_SKIP_AUTH=false; CORS_ORIGINS=https://localhost (temporary)"],
            ["Secrets", "the seven SSM names above"],
            ["Network", "same VPC as RDS; task SG = API SG that RDS already trusts"],
        ],
    )
    para(
        doc,
        "When Active, copy the HTTPS URL (example shape: https://xxxx.ecs.us-east-1.on.aws). "
        "That is $ApiUrl for web, mobile, and VITE_API_URL. Probe: curl.exe $ApiUrl/api/health"
    )

    heading(doc, "8.6 Frontend to S3", 2)
    code_block(
        doc,
        """aws s3 mb "s3://$WebBucket" --region $Region
aws s3api put-public-access-block --bucket $WebBucket --public-access-block-configuration "BlockPublicAcls=true,IgnorePublicAcls=true,BlockPublicPolicy=true,RestrictPublicBuckets=true"
cd frontend
npm ci
$env:VITE_API_URL = $ApiUrl
npm run build
aws s3 sync dist/ "s3://$WebBucket/" --delete
""",
    )
    para(
        doc,
        "VITE_* variables are compile-time. If the API hostname changes, you must rebuild. "
        "Firebase web keys stay in frontend/.env locally and in GitHub secrets for CI."
    )

    heading(doc, "8.7 CloudFront", 2)
    numbered(
        doc,
        [
            "Create OAC (sign requests, origin type S3).",
            "Create distribution: S3 origin, OAC, redirect HTTP→HTTPS, default root object index.html.",
            "Paste the console’s bucket policy onto the bucket.",
            "Custom error responses: HTTP 403 and 404 → /index.html with response code 200.",
            "Wait until Enabled. $WebUrl = https://dxxxx.cloudfront.net",
            "Set ECS env CORS_ORIGINS to exactly that origin; redeploy the service.",
            "Firebase Authentication → Settings → Authorized domains → add dxxxx.cloudfront.net (host only).",
        ],
    )

    heading(doc, "8.8 Verify", 2)
    numbered(
        doc,
        [
            "GET $ApiUrl/api/health → {\"status\":\"ok\",\"service\":\"pulse-track\"}.",
            "Open $WebUrl, sign in, create a task, refresh.",
            "Optional: curl.exe -H \"Authorization: Bearer <firebase-id-token>\" $ApiUrl/api/users/me",
        ],
    )

    heading(doc, "9. Copying local SQLite data into RDS")
    para(
        doc,
        "Local data lives in backend/pulse_track.db. Production is PostgreSQL. They are not "
        "the same file. Use backend/scripts/copy_sqlite_to_postgres.py. It copies users, goals, "
        "tasks, and activities, preserves IDs, and resets Postgres sequences. Skip leftover "
        "SQLite-only columns (focus_score, effort_logs, …)."
    )
    para(doc, "Order of tables (foreign keys): users → goals → tasks → activities.")
    para(
        doc,
        "RDS is private, so a laptop cannot connect unless you temporarily (1) set the instance "
        "publicly accessible and (2) allow your public IP on 5432, run the script, then revert "
        "both. That is what was done for the first import in this account. Do not leave 5432 open."
    )
    code_block(
        doc,
        "cd backend\n"
        "$env:DATABASE_URL = aws ssm get-parameter --region us-east-1 --name /pulse-track/DATABASE_URL --query Parameter.Value --output text\n"
        ".venv\\Scripts\\python.exe scripts\\copy_sqlite_to_postgres.py --replace\n",
    )
    para(
        doc,
        "--replace truncates existing Postgres rows. Users are matched at runtime by firebase_uid, "
        "so sign in with the same Google account you used locally."
    )

    heading(doc, "10. CI/CD after the first deploy")
    para(
        doc,
        "Workflow: .github/workflows/deploy.yml. Trigger: push to master (backend/, frontend/, or the workflow file) or manual Run workflow."
    )
    table(
        doc,
        ["GitHub secrets", "GitHub variables"],
        [
            ["AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY", "AWS_REGION (default us-east-1)"],
            ["VITE_FIREBASE_* (six web config keys)", "ECS_CLUSTER (required)"],
            ["", "ECR_REPOSITORY, ECS_SERVICE, ECS_TASK_DEFINITION, ECS_CONTAINER_NAME"],
            ["", "S3_BUCKET, CLOUDFRONT_DISTRIBUTION_ID, VITE_API_URL"],
        ],
    )
    para(
        doc,
        "API job: docker build in backend/ → push :sha and :latest → register new task definition "
        "with the new image → ecs update-service and wait for stability. Web job runs only if "
        "S3_BUCKET is set. If RDS is stopped, the wait-for-stability step fails — start the database first."
    )

    heading(doc, "11. Updating and rolling back")
    bullets(
        doc,
        [
            "API-only change — merge to master; Actions pushes a new image. Or docker push + ecs update-service --force-new-deployment.",
            "Web-only change — rebuild with the same VITE_API_URL, s3 sync --delete, CloudFront invalidation /*.",
            "Rollback API — in ECS, update the service to a previous task definition revision.",
            "Rollback web — sync an older dist/ (or re-run CI on an older commit).",
            "Config-only (CORS, secrets) — change SSM or task env and force a new deployment; no image rebuild unless the code changed.",
        ],
        bold_lead=True,
    )

    heading(doc, "12. Inventory of this AWS account (August 2026)")
    para(
        doc,
        "Useful when you come back months later. Hostnames can rotate if you recreate Express Mode; GitHub variable VITE_API_URL must match."
    )
    table(
        doc,
        ["Resource", "This account"],
        [
            ["AWS account", "529088267227"],
            ["Region in use", "us-east-1"],
            ["RDS", "pulse-track-db.c4hce2k0ays5.us-east-1.rds.amazonaws.com (private)"],
            ["RDS SG", "sg-02903957875e30665 (pulse-track-rds-sg)"],
            ["ECR repo", "pulse-track-api"],
            ["S3 web bucket", "pulse-track-web-529088267227"],
            ["ECS", "cluster default, service pulse-track-api (Express Mode / Fargate)"],
            ["Example API URL", "https://pu-33c5978573c14366842f5c7087cb4002.ecs.us-east-1.on.aws"],
            ["SSM prefix", "/pulse-track/*"],
        ],
    )

    heading(doc, "13. Alternatives we considered (interview gold)")
    table(
        doc,
        ["Option", "Why we did not use it as the production default"],
        [
            ["Single EC2 + SQLite", "No sharing with mobile, no second API replica, you patch the OS, disk is a single point of failure."],
            ["Serve React from FastAPI", "Couples UI deploys to API deploys; worse caching; still need HTTPS and a process manager."],
            ["App Runner", "Best original fit; closed to new customers 30 Apr 2026. Existing customers can keep using it."],
            ["Elastic Beanstalk", "Works, but Express Mode maps more directly to “container in, HTTPS URL out”."],
            ["Lambda + API Gateway", "Would need a different packaging story (Mangum, cold starts, connection pooling to RDS). Fine later, not the current Dockerfile."],
            ["Amazon Cognito", "Native AWS identity. Firebase was already implemented and is cross-platform. Switching is a project, not a checkbox."],
            ["Public S3 website hosting", "No easy custom error HTTPS SPA setup compared with CloudFront+OAC; bucket becomes a public origin."],
            ["Secrets in the image", "Anyone with ECR pull permission gets the database password. Rotation requires a rebuild."],
        ],
    )

    heading(doc, "14. Interview question bank")
    para(
        doc,
        "Practice out loud. Keep answers to 45–90 seconds unless they ask you to go deeper."
    )
    qa(
        doc,
        "Walk me through the architecture.",
        "Static React app on S3, HTTPS via CloudFront with OAC. API is a FastAPI container on ECS Fargate behind an ALB. Data is RDS PostgreSQL in a private subnet. Clients authenticate with Firebase and send a Bearer token; the API verifies it and scopes rows by user. Secrets come from SSM. GitHub Actions builds the image and the SPA.",
    )
    qa(
        doc,
        "Why not put the database on the same container filesystem?",
        "Containers are ephemeral. A new deploy would wipe SQLite. Two tasks cannot share a SQLite file safely. RDS outlives tasks and serves web and mobile together.",
    )
    qa(
        doc,
        "How does a user stay logged in across web and a future mobile app?",
        "Same Firebase project. Each client obtains an ID token after sign-in. The API does not care which client it is — only that the token is valid and the firebase_uid maps to a users row.",
    )
    qa(
        doc,
        "What is CORS doing here?",
        "The browser refuses to let JavaScript on the CloudFront origin call a different API origin unless the API lists that origin. Native mobile does not use CORS. Getting the CloudFront URL wrong in CORS_ORIGINS is a common production bug.",
    )
    qa(
        doc,
        "How do you keep S3 private but still serve the website?",
        "Block Public Access on the bucket. CloudFront uses Origin Access Control so only that distribution can GetObject. Users never see the S3 website endpoint.",
    )
    qa(
        doc,
        "Why Fargate instead of EC2 launch type?",
        "No capacity to patch, no ASG of instances for a single small API. We pay per vCPU/memory of the task. If we later need GPUs or bursty custom AMIs, we could move the same task definition to EC2.",
    )
    qa(
        doc,
        "How do you deploy without downtime?",
        "ECS starts a new task from the new image, health-checks /api/health, then drains the old task. ALB only sends traffic to healthy targets. The SPA is a separate deploy; CloudFront may serve a cached index.html until invalidation.",
    )
    qa(
        doc,
        "Where would this design break at 10× traffic?",
        "db.t4g.micro becomes the bottleneck before Fargate does. Next steps: RDS larger class or read replica for analytics, raise ECS max tasks, add CloudWatch alarms on ALB 5xx and RDS CPU. The API is already stateless so horizontal scale is straightforward.",
    )
    qa(
        doc,
        "How are secrets rotated?",
        "Update the SSM parameter, then force a new ECS deployment so tasks fetch the new value. Change RDS password in both RDS and DATABASE_URL. Do not rebuild the image unless code changed.",
    )
    qa(
        doc,
        "What did you do when App Runner stopped accepting new customers?",
        "Switched to Amazon ECS Express Mode, which AWS documents as the App Runner replacement. Same Dockerfile, same SSM parameters, same RDS. We traded a proprietary App Runner URL for an ECS service plus ALB that CI can update with the official GitHub Actions.",
    )
    qa(
        doc,
        "How is tenant isolation implemented?",
        "Not separate databases. One schema, every query tied to current_user.id after token verify. firebase_uid is unique. That is a shared-schema multi-tenant model appropriate for a personal productivity app.",
    )
    qa(
        doc,
        "Why is VITE_API_URL a build-time variable?",
        "Vite inlines import.meta.env.VITE_* into the JavaScript bundle. There is no server to read runtime env for the SPA. Changing the API hostname means a new frontend build and S3 sync.",
    )

    heading(doc, "15. Troubleshooting cheatsheet")
    table(
        doc,
        ["Symptom", "Likely cause"],
        [
            ["ALB/ECS unhealthy", "RDS SG, wrong DATABASE_URL, Firebase private key missing \\n, task not in RDS VPC"],
            ["Firebase Auth is not configured", "SSM names ≠ env names; execution role missing ssm:GetParameters"],
            ["Chrome CORS error", "CORS_ORIGINS not exactly the CloudFront https origin; service not redeployed"],
            ["Google sign-in blocked", "CloudFront host missing from Firebase authorized domains"],
            ["Refresh on /board is 403/404", "CloudFront custom errors not set to /index.html 200"],
            ["UI still calls localhost", "Forgot VITE_API_URL at build time"],
            ["Empty board after login on AWS", "New firebase_uid row; local data not copied, or copied under a different uid"],
            ["GitHub Action never runs", "Push was not to master"],
            ["iam:PassRole denied", "GitHub IAM policy missing the ECS execution/task role ARNs"],
            ["Deploy waits then fails", "RDS stopped, or secrets cannot be pulled"],
        ],
    )

    heading(doc, "16. Cost (order of magnitude)")
    para(
        doc,
        "Always check the AWS pricing page for the region. Ballpark for light personal use: "
        "RDS db.t4g.micro often dominates ($10–25/month). Fargate + ALB is the next line. "
        "S3, CloudFront, ECR, SSM are usually cents to a few dollars. Stop ECS desired count "
        "and stop RDS if you are not using the environment — data remains if you stop rather than delete RDS."
    )

    heading(doc, "17. First-hour checklist on a new account")
    numbered(
        doc,
        [
            "CLI + Docker + Firebase project.",
            "VPC noted; two SGs; RDS private PostgreSQL.",
            "Seven SSM parameters.",
            "ECR push of backend/Dockerfile.",
            "ECS Express service, health green.",
            "S3 bucket private; npm run build with VITE_API_URL; sync.",
            "CloudFront + OAC + SPA errors; CORS + Firebase domain.",
            "Optional: copy_sqlite_to_postgres.py then lock RDS down.",
            "GitHub secrets/variables; merge to master; watch Actions.",
        ],
    )

    heading(doc, "18. Related files in the repo")
    table(
        doc,
        ["Path", "Role"],
        [
            ["docs/DEPLOY-AWS.md", "Shorter command-oriented notes (may still mention App Runner in older sections)"],
            ["backend/Dockerfile", "API image"],
            ["backend/.dockerignore", "Keeps secrets out of the image"],
            ["backend/app/config.py", "Env + postgres:// rewrite to psycopg2"],
            ["backend/app/auth.py", "Firebase Admin verify"],
            ["backend/app/main.py", "CORS, routers, create_all, SQLite-only migrations"],
            ["backend/scripts/copy_sqlite_to_postgres.py", "Local → RDS copy"],
            [".github/workflows/deploy.yml", "master → ECR + ECS (+ S3/CloudFront)"],
            ["frontend/src/api.js", "fetch(VITE_API_URL + path) with Bearer token"],
        ],
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(
        "End of document. Rebuild this file after major architecture changes:  "
        "python docs/generate_deployment_docx.py"
    )
    set_run_font(r, size=10, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
